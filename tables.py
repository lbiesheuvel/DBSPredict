import numpy as np
from pylatexenc.latexencode import utf8tolatex
import scipy.stats as stats
import pandas as pd
from docx import Document


def __generate_categorical_statistics_table(table_data):
    n_outcomes = len(table_data["outcomes"])

    latex_table = f"""\\begin{{table}}[]
\\centering
\\resizebox{{\\textwidth}}{{!}}{{%
\\begin{{tabular}}{{@{{}}rlr{n_outcomes * 'r'}@{{}}}}
\\toprule
 &  & & \\multicolumn{{{n_outcomes}}}{{c}}{{Mean absolute improvement (Postoperative vs. Preoperative)}} \\\\ \\cmidrule(l){{{4}-{n_outcomes+3}}} 
 &  & n (\%)"""
    for i in range(n_outcomes):
        latex_table += f" & {utf8tolatex(table_data['outcomes'][i])}"
    latex_table += " \\\ \midrule\n"
    for i, feature_dict in enumerate(table_data["inner_data"]):
        feature_name = feature_dict["name"]
        feature_name = utf8tolatex(feature_name)
        total_feature_means = feature_dict["means"]
        feature_categories = feature_dict["categories"]
        total_sum = sum([category["count"] for category in feature_categories])

        for j, category in enumerate(feature_categories):
            category_name = category["name"]
            category_name = utf8tolatex(category_name)
            category_count = category["count"]
            category_means = category["means"]
            category_sds = category["sds"]  # Assuming 'sds' key exists
            if j == 0:
                latex_table += f"{feature_name} & {category_name} & {category_count} ({category_count/total_sum*100:.1f}\\%)"
            else:
                latex_table += f" & {category_name} & {category_count} ({category_count/total_sum*100:.1f}\\%)"
            for k in range(n_outcomes):
                mean = category_means[k]
                sd = category_sds[k]  # Retrieve the standard deviation
                latex_table += f" & {mean:.1f} ± {sd:.1f}"  # Format with mean ± SD
            latex_table += " \\\ \n"

    latex_table += """\\bottomrule
\\end{tabular}
}
\\caption{Descriptive statistics of categorical variables with mean absolute improvement and standard deviations.}
\\label{categorical_description_sd}
\\end{table}"""

    return latex_table


def generate_numerical_statistics_table(table_data):
    # Table with feature names, mean, and correlation to outcome variables
    n_outcomes = len(table_data["outcomes"])
    latex_table = f"""\\begin{{table}}[]
\\centering
\\resizebox{{\\textwidth}}{{!}}{{%
\\begin{{tabular}}{{@{{}}rr{'r' * n_outcomes}@{{}}}}
\\toprule
 &  & \\multicolumn{{{n_outcomes}}}{{c}}{{Outcome correlations (Postoperative ON DBS OFF Med)}} \\\\ \\cmidrule(l){{3-{n_outcomes+2}}}
 """
    latex_table += f" & Mean (95\% CI)"
    for i in range(n_outcomes):
        latex_table += f" & {utf8tolatex(table_data['outcomes'][i])}"
    latex_table += " \\\\ \\midrule"
    for i, feature_dict in enumerate(table_data["inner_data"]):
        feature_name = feature_dict["name"]
        feature_name = utf8tolatex(feature_name)
        feature_mean = feature_dict["mean"]
        feature_std = feature_dict["std"]
        feature_n = feature_dict["n"]
        feature_correlations = feature_dict["correlations"]
        # Calculate standard deviation of the mean
        std_str = f"± {feature_std:.1f}"
        # ci = stats.t.interval(0.95, df=feature_n-1, loc=feature_mean, scale=feature_std/np.sqrt(feature_n))
        # ci_str = f"({ci[0]:.1f}-{ci[1]:.1f})"
        latex_table += f"\n{feature_name} & {feature_mean:.1f} {std_str}"
        for j in range(n_outcomes):
            correlation = feature_correlations[j]
            latex_table += f" & {correlation:.2f}"
        latex_table += " \\\\"
    latex_table += """\\bottomrule
\\end{tabular}
}
\\caption{Descriptive statistics of numerical variables.}
\\label{numerical_description}
\\end{table}"""

    return latex_table


# We create some example data
def numerical_statistics_table(data, candidate_features):
    categorical_features = ["Sex", "Impulse control disorder"]
    numerical_features = [
        feature for feature in candidate_features if feature not in categorical_features
    ]
    outcomes = [
        "UPDRS-III Total",
        "UPDRS-III Tremor",
        "UPDRS-III Axial",
        "UPDRS-III Rigidity + Bradykinesia",
    ]
    table_data = {}
    table_data["outcomes"] = outcomes
    table_data["inner_data"] = []
    for feature in numerical_features:
        feature_dict = {}
        feature_dict["name"] = feature
        feature_dict["n"] = data[feature].count()
        feature_dict["mean"] = data[feature].mean()
        feature_dict["std"] = data[feature].std()
        feature_dict["correlations"] = [
            data[feature].corr(
                data["Postoperative total UPDRS-III OFF medication ON DBS score"]
            ),
            data[feature].corr(
                data[
                    "Postoperative total tremor OFF medication ON DBS score (UPDRS-III subscore)"
                ]
            ),
            data[feature].corr(
                data[
                    "Postoperative axial OFF medication ON DBS score (UPDRS-III subscore)"
                ]
            ),
            data[feature].corr(
                data[
                    "Postoperative total bradykinesia + rigidity OFF medication ON DBS score (UPDRS-III subscore)"
                ]
            ),
        ]
        table_data["inner_data"].append(feature_dict)

    return generate_numerical_statistics_table(
        table_data
    ), generate_numerical_statistics_dataframe(table_data)


def generate_numerical_statistics_dataframe(table_data):
    # Initialize a list to hold row data
    rows = []

    # Create column names for the DataFrame
    columns = ["Feature", "Mean ± SD"] + [
        f"Correlation with {outcome}" for outcome in table_data["outcomes"]
    ]

    # Iterate over each feature to collect its statistics
    for feature_dict in table_data["inner_data"]:
        # Extract feature statistics
        feature_name = feature_dict["name"]
        feature_mean = feature_dict["mean"]
        feature_std = feature_dict["std"]
        mean_std_str = f"{feature_mean:.1f} ± {feature_std:.1f}"

        # Collect correlations for each outcome
        correlations = [
            f"{feature_dict['correlations'][i]:.2f}"
            for i in range(len(table_data["outcomes"]))
        ]

        # Create a row for the current feature
        row = [feature_name, mean_std_str] + correlations
        rows.append(row)

    # Create the DataFrame
    df = pd.DataFrame(rows, columns=columns)

    return df


def generate_categorical_statistics_df(table_data):
    rows = []
    n_outcomes = len(table_data["outcomes"])
    outcomes = table_data["outcomes"]

    # Header for DataFrame
    columns = ["Feature", "Category", "n (%)"] + [
        f"{outcome} Mean ± SD" for outcome in outcomes
    ]

    for feature_dict in table_data["inner_data"]:
        feature_name = feature_dict["name"]
        total_feature_means = feature_dict["means"]
        feature_categories = feature_dict["categories"]
        total_sum = sum([category["count"] for category in feature_categories])

        for category in feature_categories:
            category_name = category["name"]
            category_count = category["count"]
            category_percentage = category_count / total_sum * 100
            category_means = category["means"]
            category_sds = category["sds"]
            row = [
                feature_name,
                category_name,
                f"{category_count} ({category_percentage:.1f}%)",
            ]

            # Append means and SDs for each outcome
            for mean, sd in zip(category_means, category_sds):
                row.append(f"{mean:.1f} ± {sd:.1f}")
            rows.append(row)

            # To avoid repeating the feature name for subsequent categories
            feature_name = ""

    return pd.DataFrame(rows, columns=columns)


def categorical_statistics_table(data):
    category_dict = {
        "Sex": {1: "Male", 2: "Female"},
        "Impulse control disorder": {1: "Yes", 2: "No"},
    }

    categorical_features = ["Sex", "Impulse control disorder"]
    outcomes = [
        "UPDRS-III Total",
        "UPDRS-III Tremor",
        "UPDRS-III Axial",
        "UPDRS-III Rigidity + Bradykinesia",
    ]
    table_data = {}
    table_data["outcomes"] = outcomes
    table_data["inner_data"] = []
    for i, feature in enumerate(categorical_features):
        feature_dict = {}
        feature_dict["name"] = feature
        y_feature_total = data[
            "Postoperative total UPDRS-III OFF medication ON DBS score"
        ]  # 3 more needed
        y_feature_rigidity_total = data[
            "Postoperative total bradykinesia + rigidity OFF medication ON DBS score (UPDRS-III subscore)"
        ]
        y_feature_tremor_total = data[
            "Postoperative total tremor OFF medication ON DBS score (UPDRS-III subscore)"
        ]
        y_feature_axial_total = data[
            "Postoperative axial OFF medication ON DBS score (UPDRS-III subscore)"
        ]

        feature_dict["means"] = [
            y_feature_total.mean(),
            y_feature_rigidity_total.mean(),
            y_feature_tremor_total.mean(),
            y_feature_axial_total.mean(),
        ]
        feature_dict["categories"] = []
        categorical_feature_value_counts = data[feature].value_counts()
        for category in categorical_feature_value_counts.index:
            category_name = category_dict[feature][category]
            category_count = categorical_feature_value_counts[category]
            relevant_rows = data[data[feature] == category]
            # remove rows where the outcome or baseline feature is missing
            # relevant_rows_total = relevant_rows.dropna(subset=['Postoperative total UPDRS-III OFF medication ON DBS score', 'Total UPDRS-III OFF score'])
            # relevant_rows_rigidity = relevant_rows.dropna(subset=['Postoperative total bradykinesia + rigidity OFF medication ON DBS score (UPDRS-III subscore)', 'Total bradykinesia + rigidity OFF score (UPDRS-III subscore)'])
            # relevant_rows_tremor = relevant_rows.dropna(subset=['Postoperative total tremor OFF medication ON DBS score (UPDRS-III subscore)', 'Total tremor OFF score (UPDRS-III subscore)'])
            # relevant_rows_axial = relevant_rows.dropna(subset=['Postoperative axial OFF medication ON DBS score (UPDRS-III subscore)', 'Axial OFF score (UPDRS-III subscore)'])

            pre_operative_feature_category = relevant_rows["Total UPDRS-III OFF score"]
            postoperative_feature_category = relevant_rows[
                "Postoperative total UPDRS-III OFF medication ON DBS score"
            ]
            pre_operative_feature_category_feature_rigidity_category = relevant_rows[
                "Total bradykinesia + rigidity OFF score (UPDRS-III subscore)"
            ]
            postoperative_feature_category_feature_rigidity_category = relevant_rows[
                "Postoperative total bradykinesia + rigidity OFF medication ON DBS score (UPDRS-III subscore)"
            ]
            pre_operative_feature_category_feature_tremor_category = relevant_rows[
                "Total tremor OFF score (UPDRS-III subscore)"
            ]
            postoperative_feature_category_feature_tremor_category = relevant_rows[
                "Postoperative total tremor OFF medication ON DBS score (UPDRS-III subscore)"
            ]
            pre_operative_feature_category_feature_axial_category = relevant_rows[
                "Axial OFF score (UPDRS-III subscore)"
            ]
            postoperative_feature_category_feature_axial_category = relevant_rows[
                "Postoperative axial OFF medication ON DBS score (UPDRS-III subscore)"
            ]
            improvement_feature_category = (
                pre_operative_feature_category - postoperative_feature_category
            )
            improvement_feature_rigidity_category = (
                pre_operative_feature_category_feature_rigidity_category
                - postoperative_feature_category_feature_rigidity_category
            )
            improvement_feature_tremor_category = (
                pre_operative_feature_category_feature_tremor_category
                - postoperative_feature_category_feature_tremor_category
            )
            improvement_feature_axial_category = (
                pre_operative_feature_category_feature_axial_category
                - postoperative_feature_category_feature_axial_category
            )
            # We calculate the 95% confidence interval for the improvements
            improvement_feature_category_std = improvement_feature_category.std()
            improvement_feature_rigidity_category_std = (
                improvement_feature_rigidity_category.std()
            )
            improvement_feature_tremor_category_std = (
                improvement_feature_tremor_category.std()
            )
            improvement_feature_axial_category_std = (
                improvement_feature_axial_category.std()
            )
            improvement_feature_category_mean = improvement_feature_category.mean()
            improvement_feature_rigidity_category_mean = (
                improvement_feature_rigidity_category.mean()
            )
            improvement_feature_tremor_category_mean = (
                improvement_feature_tremor_category.mean()
            )
            improvement_feature_axial_category_mean = (
                improvement_feature_axial_category.mean()
            )
            improvement_feature_category_n = improvement_feature_category.count()
            improvement_feature_rigidity_category_n = (
                improvement_feature_rigidity_category.count()
            )
            improvement_feature_tremor_category_n = (
                improvement_feature_tremor_category.count()
            )
            improvement_feature_axial_category_n = (
                improvement_feature_axial_category.count()
            )
            improvement_feature_category_se = (
                improvement_feature_category_std
                / np.sqrt(improvement_feature_category_n)
            )
            improvement_feature_rigidity_category_se = (
                improvement_feature_rigidity_category_std
                / np.sqrt(improvement_feature_rigidity_category_n)
            )
            improvement_feature_tremor_category_se = (
                improvement_feature_tremor_category_std
                / np.sqrt(improvement_feature_tremor_category_n)
            )
            improvement_feature_axial_category_se = (
                improvement_feature_axial_category_std
                / np.sqrt(improvement_feature_axial_category_n)
            )
            improvement_feature_category_lower_bound = (
                improvement_feature_category_mean
                - 1.96 * improvement_feature_category_se
            )
            improvement_feature_rigidity_category_lower_bound = (
                improvement_feature_rigidity_category_mean
                - 1.96 * improvement_feature_rigidity_category_se
            )
            improvement_feature_tremor_category_lower_bound = (
                improvement_feature_tremor_category_mean
                - 1.96 * improvement_feature_tremor_category_se
            )
            improvement_feature_axial_category_lower_bound = (
                improvement_feature_axial_category_mean
                - 1.96 * improvement_feature_axial_category_se
            )
            improvement_feature_category_upper_bound = (
                improvement_feature_category_mean
                + 1.96 * improvement_feature_category_se
            )
            improvement_feature_rigidity_category_upper_bound = (
                improvement_feature_rigidity_category_mean
                + 1.96 * improvement_feature_rigidity_category_se
            )
            improvement_feature_tremor_category_upper_bound = (
                improvement_feature_tremor_category_mean
                + 1.96 * improvement_feature_tremor_category_se
            )
            improvement_feature_axial_category_upper_bound = (
                improvement_feature_axial_category_mean
                + 1.96 * improvement_feature_axial_category_se
            )
            lower_bounds = [
                improvement_feature_category_lower_bound,
                improvement_feature_tremor_category_lower_bound,
                improvement_feature_axial_category_lower_bound,
                improvement_feature_rigidity_category_lower_bound,
            ]
            upper_bounds = [
                improvement_feature_category_upper_bound,
                improvement_feature_tremor_category_upper_bound,
                improvement_feature_axial_category_upper_bound,
                improvement_feature_rigidity_category_upper_bound,
            ]

            category_means = [
                improvement_feature_category.mean(),
                improvement_feature_tremor_category.mean(),
                improvement_feature_axial_category.mean(),
                improvement_feature_rigidity_category.mean(),
            ]

            # we calculate the std of the mean improvement
            sds = [
                improvement_feature_category_std,
                improvement_feature_tremor_category_std,
                improvement_feature_axial_category_std,
                improvement_feature_rigidity_category_std,
            ]

            feature_dict["categories"].append(
                {
                    "name": category_name,
                    "count": category_count,
                    "means": category_means,
                    "lower_bounds": lower_bounds,
                    "upper_bounds": upper_bounds,
                    "sds": sds,
                }
            )

        table_data["inner_data"].append(feature_dict)
    return (
        __generate_categorical_statistics_table(table_data),
        generate_categorical_statistics_df(table_data),
    )


def generate_results_table(results):
    latex_table = "\\begin{table}[]\n\\resizebox{\\textwidth}{!}{\n\\begin{tabular}{@{}lrrrr@{}}\n\\toprule\n"
    latex_table += "Model & Number of training samples & Number of testing samples & Best model type & RMSE (Testing set) \\\\ \\midrule\n"

    for result in results:
        latex_table += f"{result['name']} & {result['n_records_training_data']} & {result['n_records_testing_data']} & {result['model_type']} & {result['rmse']:.1f} ({result['lower bound']:.1f}-{result['upper bound']:.1f}) \\\\\n"

    latex_table += "\\bottomrule\n\\end{tabular}\n}\n\\caption{Evaluated performance of the models}\n\\label{results_table}\n\\end{table}"

    return latex_table


def generate_results_dataframe(results):
    # Create a DataFrame from the results
    df = pd.DataFrame(results)

    # Add a new column to format RMSE with its bounds
    df["RMSE (Testing set)"] = df.apply(
        lambda x: f"{x['rmse']:.1f} ({x['lower bound']:.1f}-{x['upper bound']:.1f})",
        axis=1,
    )

    # Add a column for R squared results too
    df["R squared (Testing set)"] = df.apply(
        lambda x: f"{x['r squared']:.2f} ({x['r squared lower bound']:.2f}-{x['r squared upper bound']:.2f})",
        axis=1,
    )

    # Select and rename columns to match the LaTeX table
    df = df[
        [
            "name",
            "n_records_training_data",
            "n_records_testing_data",
            "model_type",
            "RMSE (Testing set)",
            "R squared (Testing set)",
        ]
    ]
    df.columns = [
        "Model",
        "Number of training samples",
        "Number of testing samples",
        "Best model type",
        "RMSE testing set (95% CI)",
        "R squared testing set (95% CI)",
    ]

    return df


def generate_feature_table(models, all_features):
    latex_table = (
        "\\begin{table}[]\n\\resizebox{\\textwidth}{!}{\n\\begin{tabular}{l"
        + "c" * len(models)
        + "}\n\\toprule\n"
    )
    latex_table += (
        " & "
        + " & ".join([f"\\textbf{{{utf8tolatex(model['name'])}}}" for model in models])
        + " \\\\ \\midrule\n"
    )

    for i, feature in enumerate(all_features):
        row = [utf8tolatex(feature)]
        for model in models:
            if feature in model["features"]:
                row.append("X")  # or replace with "\\checkmark" for a checkmark
            else:
                row.append("")
        if i % 2 == 0:
            latex_table += "\\rowcolor{lightgray}"
        latex_table += " & ".join(row) + " \\\\\n"

    latex_table += "\\bottomrule\n\\end{tabular}\n}\n\\caption{Included features for each model}\n\\label{table_features}\n\\end{table}"

    return latex_table


def generate_feature_dataframe(models, all_features):
    # Initialize an empty dictionary to store data
    data = {}

    # Fill in the dictionary with model names as keys and lists of feature presence as values
    for model in models:
        data[model["name"]] = [
            "X" if feature in model["features"] else "" for feature in all_features
        ]

    # Convert the dictionary into a DataFrame, with all_features as the index
    df = pd.DataFrame(data, index=all_features)

    return df


def export_dataframe_to_word(df, filename):
    # Create a new Document
    doc = Document()
    # Add a title
    doc.add_heading("DataFrame Export", 0)

    # Add a table with an extra row for the headers
    table = doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])

    # Style the table for better readability
    table.style = "Table Grid"

    # Add the DataFrame headers to the table
    for j, col_name in enumerate(df.columns):
        table.cell(0, j).text = col_name

    # Populate the table with the DataFrame's data
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            table.cell(i + 1, j).text = str(value)

    # Save the document
    doc.save(f"./paper/tables/{filename}.docx")
